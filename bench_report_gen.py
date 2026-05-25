#!/usr/bin/env python3
"""
bench_report_gen.py — Build a Word (.docx) report from bench_image_lab.py output.

Usage (on Windows, after SCP):
    python bench_report_gen.py

Reads  : C:\\Temp\\bench\\bench_results.json  +  C:\\Temp\\bench\\*.png
Writes : C:\\Temp\\bench\\arthur_image_lab_benchmark_<date>.docx
"""

from __future__ import annotations
import json
import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
BENCH_DIR    = Path(r"C:\Temp\bench")
RESULTS_JSON = BENCH_DIR / "bench_results.json"
date_tag     = datetime.datetime.now().strftime("%Y-%m-%d_%H%M")
OUT_DOCX     = BENCH_DIR / f"arthur_image_lab_benchmark_{date_tag}.docx"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_col_width(table, col_idx: int, width_inches: float):
    for row in table.rows:
        row.cells[col_idx].width = int(width_inches * 914400)  # EMU


def shade_cell(cell, hex_color: str = "D9E1F2"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def bold_cell(cell, text: str, size_pt: int = 10, hex_color: str = "D9E1F2"):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = True
    run.font.size = Pt(size_pt)
    shade_cell(cell, hex_color)


def fmt_s(v) -> str:
    if v is None:
        return "—"
    return f"{float(v):.1f} s"


def fmt_sps(v) -> str:
    if v is None:
        return "—"
    return f"{float(v):.2f} it/s"


# ---------------------------------------------------------------------------
# Load results
# ---------------------------------------------------------------------------
results: list[dict] = json.loads(RESULTS_JSON.read_text())

# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------
doc = Document()

# Page margins (narrow)
for section in doc.sections:
    section.top_margin    = int(0.75 * 914400)
    section.bottom_margin = int(0.75 * 914400)
    section.left_margin   = int(1.00 * 914400)
    section.right_margin  = int(1.00 * 914400)

# ── Title ────────────────────────────────────────────────────────────────────
title_para = doc.add_heading("Arthur Image Lab — Benchmark Report", level=0)
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph(
    f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}   |   "
    f"Host: 192.168.0.87   |   GPU: RTX 5060 Ti (15.48 GB)"
)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)
sub.runs[0].font.size = Pt(9)

doc.add_paragraph()

# ── Environment ───────────────────────────────────────────────────────────────
doc.add_heading("Environment", level=1)

env_rows = [
    ("GPU",        "NVIDIA RTX 5060 Ti (15.48 GB, Blackwell SM100+)"),
    ("CUDA",       "12.8.90"),
    ("PyTorch",    "2.11.0+cu128"),
    ("torchao",    "0.17.0+cu128  (NVFP4 / MX formats)"),
    ("OS",         "Ubuntu 22.04 LTS"),
    ("Service",    "arthur-imglab.service (port 8002)"),
    ("Prompt",
     "a photorealistic mountain lake at golden hour, reflections of snow-capped "
     "peaks, misty atmosphere, dramatic clouds, cinematic lighting, 8k"),
    ("Seed",       "42 (fixed)"),
    ("Resolution", "1024 × 1024 px"),
]

env_table = doc.add_table(rows=len(env_rows), cols=2)
env_table.style = "Table Grid"
for i, (k, v) in enumerate(env_rows):
    bold_cell(env_table.rows[i].cells[0], k, hex_color="EBF3FB")
    env_table.rows[i].cells[1].text = v
    env_table.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(9)

doc.add_paragraph()

# ── Summary table ─────────────────────────────────────────────────────────────
doc.add_heading("Performance Summary", level=1)

HDR = ["Model", "Quant", "Steps", "Load (s)", "Inference (s)", "Speed (it/s)", "Total (s)", "Status"]
t = doc.add_table(rows=1 + len(results), cols=len(HDR))
t.style = "Table Grid"

for i, h in enumerate(HDR):
    bold_cell(t.rows[0].cells[i], h, hex_color="1F3864")
    t.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for ri, r in enumerate(results):
    row = t.rows[ri + 1]
    row.cells[0].text = r["label"]
    row.cells[1].text = r["quant"]
    row.cells[2].text = str(r["steps"])
    row.cells[3].text = fmt_s(r.get("load_s"))
    row.cells[4].text = fmt_s(r.get("inf_s"))
    row.cells[5].text = fmt_sps(r.get("steps_per_s"))
    row.cells[6].text = fmt_s(r.get("total_s"))
    row.cells[7].text = r["status"].upper()

    # Colour status cell
    status_cell = row.cells[7]
    if r["status"] == "ok":
        shade_cell(status_cell, "C6EFCE")   # green
        status_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x27, 0x6E, 0x27)
    else:
        shade_cell(status_cell, "FFC7CE")   # red
        status_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x9C, 0x00, 0x06)

    # Zebra striping
    if ri % 2 == 1:
        for ci in range(len(HDR) - 1):
            shade_cell(row.cells[ci], "F2F2F2")

doc.add_paragraph()

# ── Per-model sections ────────────────────────────────────────────────────────
doc.add_heading("Results per Model", level=1)

for r in results:
    doc.add_heading(r["label"], level=2)

    # Detail table
    detail_rows = [
        ("Engine key",     r["engine"]),
        ("Quantization",   r["quant"]),
        ("Steps",          str(r["steps"])),
        ("Model load",     fmt_s(r.get("load_s"))),
        ("Inference",      fmt_s(r.get("inf_s"))),
        ("Speed",          fmt_sps(r.get("steps_per_s"))),
        ("Total (wall)",   fmt_s(r.get("total_s"))),
        ("VRAM after",     f"{r.get('vram_used_after_mb', '—')} MiB" if r.get("vram_used_after_mb") else "—"),
    ]

    dt = doc.add_table(rows=len(detail_rows), cols=2)
    dt.style = "Table Grid"
    for i, (k, v) in enumerate(detail_rows):
        bold_cell(dt.rows[i].cells[0], k, hex_color="EBF3FB")
        dt.rows[i].cells[0].width = int(1.6 * 914400)
        dt.rows[i].cells[1].text = v
        dt.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    if r["status"] != "ok":
        p = doc.add_paragraph()
        run = p.add_run(f"⚠  Generation failed: {r.get('error', 'unknown error')}")
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        run.font.size = Pt(9)
    else:
        img_path = BENCH_DIR / f"{r['engine']}_{r['quant']}.png"
        if img_path.exists():
            doc.add_picture(str(img_path), width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p = doc.add_paragraph()
            p.add_run(f"(image not found: {img_path})")

    doc.add_paragraph()

# ── Observations ──────────────────────────────────────────────────────────────
doc.add_heading("Key Observations", level=1)

ok_results = [r for r in results if r["status"] == "ok"]
if ok_results:
    fastest = min(ok_results, key=lambda r: r.get("total_s") or 9999)
    doc.add_paragraph(
        f"• Fastest total time:  {fastest['label']}  →  {fmt_s(fastest.get('total_s'))}",
        style="List Bullet"
    )

    best_sps = [r for r in ok_results if r.get("steps_per_s")]
    if best_sps:
        top = max(best_sps, key=lambda r: r["steps_per_s"])
        doc.add_paragraph(
            f"• Fastest inference speed:  {top['label']}  →  {fmt_sps(top['steps_per_s'])}",
            style="List Bullet"
        )

doc.add_paragraph(
    "• NVFP4 quantization requires the GPU to be fully free during nvfp4_save.py "
    "(stop the service first); CUDA kernels are mandatory — CPU-only device_map "
    "silently saves BF16 at full model size.",
    style="List Bullet"
)
doc.add_paragraph(
    "• FLUX.2 Klein 4B uses only 4 inference steps (step-distilled); "
    "its 'total time' is not directly comparable to 28-step models.",
    style="List Bullet"
)

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
doc.save(str(OUT_DOCX))
print(f"Saved: {OUT_DOCX}")
